from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).parent
OUT_FILE = OUT_DIR / "电路图智能评价系统_当前进展与投稿要求.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
TEXT = "1F2937"
MUTED = "5B6573"
RED = "9B1C1C"
GREEN = "1F5A3A"
FONT = "Microsoft YaHei"
LATIN_FONT = "Calibri"


def set_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = LATIN_FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margin(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margin(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def write_cell(cell, text, *, bold=False, color=TEXT, size=10.5):
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 11)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    run = p.add_run(text)
    set_font(run, size=16 if level == 1 else 12.5, bold=True, color=BLUE if level == 1 else DARK_BLUE)
    return p


def add_body(doc, text, *, bold_prefix=None, color=TEXT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.16
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, size=10.8, bold=True, color=color)
        r = p.add_run(text[len(bold_prefix):])
        set_font(r, size=10.8, color=color)
    else:
        r = p.add_run(text)
        set_font(r, size=10.8, color=color)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_font(r, size=10.6, color=TEXT)
    return p


def add_note(doc, label, text, fill=LIGHT_GRAY, label_color=DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.14
    label_run = p.add_run(label + " ")
    set_font(label_run, size=10.6, bold=True, color=label_color)
    body_run = p.add_run(text)
    set_font(body_run, size=10.6, color=TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_status_table(doc):
    headers = ["模块", "当前已完成", "可用于投稿的状态"]
    rows = [
        ("数据与模型", "50 份连线真值、50 份检测结果；检测模型与 CRNN OCR 权重已就绪。", "可用，但须固定数据划分并补充 OCR 独立指标。"),
        ("元件检测", "最佳记录：Precision 85.1%，Recall 68.6%，mAP@50 75.3%，mAP@50-95 57.0%。", "可作为结果表中的独立模块指标。"),
        ("连线重建", "已有骨架、CCL、P2J/JJ、LOS、近端口等路径；历史完整方案记录：PC 30.1%、GA 14.0%、CNA 11.3%。", "尚不可投稿：需全量复测，并显著改善端口/连通组指标。"),
        ("实验工程", "已加入独立输出目录、配置快照、Git 版本记录、真实消融配置；实验禁用 LLM 与附属图片/文本写入。", "已具备可复现基础，仍需生成正式全量结果包。"),
        ("测试", "视觉相关与 phase2 非 LLM 测试 18/18 通过；全量测试中 4 个 phase1 用例因未配置 LLM 凭证失败。", "视觉论文可不使用 LLM；投稿前应将 LLM 测试标注为外部依赖或隔离。"),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [1600, 4200, 3560])
    for cell, text in zip(table.rows[0].cells, headers):
        shade_cell(cell, LIGHT_BLUE)
        write_cell(cell, text, bold=True, color=DARK_BLUE, size=10.5)
    for item, completed, status in rows:
        cells = table.add_row().cells
        write_cell(cells[0], item, bold=True, color=DARK_BLUE)
        write_cell(cells[1], completed)
        write_cell(cells[2], status, color=GREEN if "可用" in status or "已具备" in status else RED)


def add_submission_table(doc):
    headers = ["投稿前必须完成", "EI 期刊最低要求", "SCI/SCIE 进一步要求"]
    rows = [
        ("贡献定义", "给出一个可验证的核心方法：元件遮蔽导线证据、尺度一致骨架重建与拓扑约束图组装。", "再增加可泛化的算法贡献，例如端子关键点预测或全局候选边优化。"),
        ("主实验", "锁定测试集，完成 50 张真值图或说明剔除原因；不得以测试集调参。", "增加跨绘图者、拍摄角度、光照或外部数据条件验证。"),
        ("指标", "检测 mAP/P/R；OCR CER 与整串准确率；连线 P/R/F1、FP/FN、GA、CNA、耗时。", "增加置信区间、统计显著性和复杂度/速度分析。"),
        ("对比与消融", "比较原始规则、CCL、完整方法；逐一去掉骨架、遮蔽、NN 过滤、LOS 等。", "与公开方法或可信复现实验做充分比较，并做误差分层分析。"),
        ("可复现性", "发布/说明固定划分、随机种子、命令、依赖版本、模型获得方式与失败案例。", "最好提供代码、权重或可申请访问的最小复现实验包。"),
        ("论文呈现", "英文摘要、方法图、伪代码、实验表、消融图、局限性与伦理/数据许可说明。", "强化问题重要性、理论/方法差异和国际相关工作覆盖。"),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [1700, 3830, 3830])
    for cell, text in zip(table.rows[0].cells, headers):
        shade_cell(cell, LIGHT_BLUE)
        write_cell(cell, text, bold=True, color=DARK_BLUE, size=10.3)
    for requirement, ei, sci in rows:
        cells = table.add_row().cells
        write_cell(cells[0], requirement, bold=True, color=DARK_BLUE)
        write_cell(cells[1], ei)
        write_cell(cells[2], sci)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.8)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    set_font(header.add_run("电路图智能评价系统 | 投稿准备技术简报"), size=8.8, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    set_font(footer.add_run("内部工作文档 | 2026-08-01"), size=8.5, color=MUTED)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(16)
    title.paragraph_format.space_after = Pt(4)
    set_font(title.add_run("电路图智能评价系统"), size=23, bold=True, color=DARK_BLUE)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(14)
    set_font(subtitle.add_run("当前已完成内容与 EI / SCI 投稿要求"), size=14, bold=True, color=BLUE)
    metadata = doc.add_paragraph()
    metadata.paragraph_format.space_before = Pt(0)
    metadata.paragraph_format.space_after = Pt(12)
    set_font(metadata.add_run("编制日期：2026-08-01  |  范围：仅基于代码、模型、实验产物与当前验证状态；不引用发表路线图。"), size=9.3, color=MUTED)

    add_note(doc, "结论：", "项目已具备完整视觉管线、数据与可复现实验基础，但当前端到端连线指标与全量复测证据不足，尚不宜直接投稿。建议先达到 EI 期刊的实证要求，再以更强的算法贡献和跨条件验证冲击 SCI/SCIE。")

    add_heading(doc, "一、目前已完成的内容")
    add_status_table(doc)

    add_heading(doc, "二、已完成的最新视觉改造")
    add_bullet(doc, "实验配置全部强制跳过 LLM；实验运行不会触发外部模型接口。")
    add_bullet(doc, "实验结果写入独立目录，并保存完整配置与 Git 版本；避免覆盖历史 CSV、图表和 benchmark 旁的结果文件。")
    add_bullet(doc, "新增元件内部遮蔽导线证据图：CCL 与骨架路径可共享该证据，并提供 w/o_Component_Mask 真实消融。")
    add_bullet(doc, "已保证遮蔽与未遮蔽骨架路径的阈值和缩放顺序一致，避免将预处理差异误解释为算法增益。")
    add_bullet(doc, "视觉相关与非 LLM 安全评测共 18 项测试通过；单图无 LLM 冒烟运行成功，且未生成附属标注图或文本。")
    add_note(doc, "重要限制：", "全量视觉基准在本轮 CPU 运行中耗时过长而被安全停止，未生成新的汇总 CSV。因此，不能将上述改造写成“已提升至某一新准确率”；投稿只能使用重新完成后的全量结果。", fill="FFF4F4", label_color=RED)

    add_heading(doc, "三、EI 期刊与 SCI/SCIE 投稿前必须完成的工作")
    add_submission_table(doc)

    add_heading(doc, "四、建议的投稿门槛与当前判断")
    add_body(doc, "EI 期刊：完成固定测试集的全量无 LLM 实验、独立 OCR 指标、可信消融、对比基线和可复现实验包后，可开始匹配工程视觉、智能制造或电子设计自动化方向期刊。")
    add_body(doc, "SCI/SCIE：除 EI 的全部要求外，当前“遮蔽 + 规则优化”还不够。需要形成一个更强、可泛化的核心贡献，例如端子关键点预测、实例分割导线掩码，或带电路约束的全局图优化，并在跨绘图者/跨采集条件上验证。")
    add_body(doc, "当前优先级：先将端口连线 F1、GA、CNA 与 FP/FN 的新旧对比完整跑出并稳定复现；若端口正确率仍接近历史 30.1%，应继续优化连线重建，不建议投稿。")

    add_heading(doc, "五、选刊与合规核验")
    add_bullet(doc, "EI Compendex 与 SCI/SCIE 是不同检索体系；不能仅凭期刊网站宣传判断收录状态，应在投稿前逐本核验。")
    add_bullet(doc, "EI 相关期刊应确认同行评审、英文摘要、规律出版、出版伦理与数字保存政策，并确认期刊范围与工程主题匹配。")
    add_bullet(doc, "SCI/SCIE 期刊需确认其当前 Web of Science 状态、期刊范围、开放获取费用、数据/代码政策、版面与篇幅要求。")
    add_bullet(doc, "论文中必须如实写出数据许可、人工标注范围、失败案例和局限性；不应对未完成的全量实验作性能承诺。")

    add_heading(doc, "六、参考依据")
    sources = [
        "Elsevier Engineering Village. Content policy and selection criteria for Compendex journals: https://www.elsevier.com/products/engineering-village/databases/selection-criteria",
        "Clarivate. Editorial selection process for the Web of Science Core Collection: https://clarivate.com/products/scientific-and-academic-research/research-discovery-and-workflow-solutions/web-ofscience-platform/web-of-science-core-collection/editorial-selection-process/editorial-selection-process/",
        "Bayer et al. Modular Graph Extraction for Handwritten Circuit Diagram Images (2024): https://arxiv.org/abs/2402.11093",
        "Bayer et al. Instance Segmentation Based Graph Extraction for Handwritten Circuit Diagram Images (2023): https://arxiv.org/abs/2301.03155",
    ]
    for source in sources:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.08
        set_font(p.add_run(source), size=8.8, color=MUTED)

    doc.core_properties.title = "电路图智能评价系统：当前已完成内容与 EI/SCI 投稿要求"
    doc.core_properties.subject = "项目状态与投稿准备清单"
    doc.core_properties.author = "Codex"
    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    main()
